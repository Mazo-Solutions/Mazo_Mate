import streamlit as st
import pandas as pd
import os
from dotenv import load_dotenv
import google.generativeai as genai
import io
from docx import Document

# Load environment variables from .env file
load_dotenv()

# Configure Gemini API key
genai.configure(api_key=os.getenv("GEMINI_API_KEY"))

# Function to generate interview questions using Gemini
def generate_interview_questions(domain, experience_level, complexity, num_questions):
    """Generate interview questions and answers using Gemini."""
    try:
        # Define generation configuration
        generation_config = {
            "temperature": 1,
            "top_p": 0.95,
            "top_k": 40,
            "max_output_tokens": 8192,
            "response_mime_type": "text/plain",
        }

        # Initialize the Gemini model
        model = genai.GenerativeModel(
            model_name="gemini-1.5-pro",
            generation_config=generation_config,
        )

        # Start a chat session
        chat_session = model.start_chat()

        # Build the prompt dynamically based on the input domain
        prompt = (
            f"Generate {num_questions} interview questions and answers for a {domain} professional "
            f"with {experience_level} years of experience. The questions should be of {complexity} complexity. "
            f"Each question should be followed by a corresponding answer."
        )

        # Send the prompt and get the response
        response = chat_session.send_message(prompt)
        return response.text.strip()
    except Exception as e:
        st.error(f"Error fetching questions: {e}")
        return ""

# Function to export data to an Excel file and return the file as a BytesIO object
def export_to_excel(data):
    """Exports the data to an Excel file and returns the file in memory as a BytesIO object."""
    try:
        # Create a pandas DataFrame from the list of dictionaries
        df = pd.DataFrame(data)

        # Save to a BytesIO buffer using openpyxl (default engine for pandas)
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, index=False, sheet_name='Interview Questions')

        output.seek(0)  # Rewind the buffer
        return output
    except Exception as e:
        st.error(f"Error exporting to Excel: {e}")
        return None

# Function to export data to a Word document and return the file as a BytesIO object
def export_to_word(data):
    """Exports the data to a Word document and returns the file in memory as a BytesIO object."""
    try:
        # Create a Word document
        doc = Document()
        doc.add_heading('Interview Questions and Answers', 0)

        for qa in data:
            # Add question and answer to Word
            doc.add_heading(qa['Question'], level=1)
            doc.add_paragraph(qa['Answer'])

        # Save to a BytesIO buffer
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)  # Rewind the buffer
        return output
    except Exception as e:
        st.error(f"Error exporting to Word: {e}")
        return None

# Main Streamlit app
def main():
    # Centered logo
    st.markdown("""
        <style>
            .center-image {
                display: flex;
                justify-content: center;
                align-items: center;
                margin-bottom: 20px;
            }
        </style>
        <div class="center-image">
            <img src="https://mazobeam.com/wp-content/uploads/2023/12/mazoid-1.png" alt="MazoBot Logo" width="200"/>
        </div>
    """, unsafe_allow_html=True)

    # Display the title "MazoMate"
    st.title("MazoMate - Interview Question Generator")

    # Configuration inputs
    domain = st.text_input("Enter a Programming/Area (e.g., Python, Java, C++, HR, Data Science, Marketing)")
    experience_level = st.number_input("Experience Level (years)", min_value=1, max_value=50, step=1, value=10)
    complexity = st.radio("Select Question Complexity", ["Basic", "Intermediate", "Advanced"])
    num_questions = st.number_input("Number of Questions to Generate", min_value=1, max_value=100, step=1, value=10)

    # Generate Questions Button
    if st.button("Generate Questions"):
        st.info("Generating interview questions. Please wait...")
        generated_content = generate_interview_questions(domain, experience_level, complexity, num_questions)

        if generated_content:
            st.success("Questions generated successfully!")
            st.write("Generated Questions and Answers")
            st.text_area("Questions & Answers", generated_content, height=300)

            # Parse questions and answers
            qa_pairs = []
            lines = generated_content.split('\n')
            for i in range(0, len(lines), 2):  # Assuming questions and answers alternate
                question = lines[i].strip() if i < len(lines) else ""
                answer = lines[i + 1].strip() if (i + 1) < len(lines) else ""
                qa_pairs.append({"Question": question, "Answer": answer})

            # Export to Excel
            excel_file = export_to_excel(qa_pairs)
            if excel_file:
                st.download_button(
                    label="Download as Excel",
                    data=excel_file,
                    file_name="Mazo_Interview_Questions.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                )

            # Export to Word
            word_file = export_to_word(qa_pairs)
            if word_file:
                st.download_button(
                    label="Download as Word",
                    data=word_file,
                    file_name="Mazo_Interview_Questions.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

if __name__ == "__main__":
    main()
